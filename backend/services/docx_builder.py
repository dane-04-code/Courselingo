"""DOCX builder service — replaces text in-place to preserve all formatting."""

from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


# ── Script detection & font mapping ──────────────────────────────────
# DOCX uses w:rFonts/@w:eastAsia for CJK text.  We set the font name to a
# well-known NotoSans family so viewers that have it installed render
# correctly; viewers without it fall back to their OS CJK font.
_SCRIPT_EAST_ASIA_FONT: dict[str, str] = {
    "jp": "Noto Sans JP",
    "kr": "Noto Sans KR",
    "sc": "Noto Sans SC",
}


def _detect_cjk_script(text: str) -> str | None:
    """Return a script key if *text* contains CJK characters, else None."""
    for char in text:
        cp = ord(char)
        if 0x3040 <= cp <= 0x30FF:                              # Hiragana / Katakana
            return "jp"
        if 0xAC00 <= cp <= 0xD7AF or 0x1100 <= cp <= 0x11FF:  # Hangul
            return "kr"
        if (0x4E00 <= cp <= 0x9FFF                             # CJK Unified Ideographs
                or 0x3400 <= cp <= 0x4DBF                      # CJK Extension A
                or 0x20000 <= cp <= 0x2A6DF):                  # CJK Extension B
            return "sc"
    return None


def _apply_east_asia_font(run, font_name: str) -> None:
    """Set w:rFonts/@w:eastAsia on *run* so CJK glyphs use the right font."""
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def _replace_paragraph_text(paragraph, new_text: str) -> None:
    """Replace the visible text of *paragraph* while preserving first-run formatting.

    Also sets the east-Asian font on the first run when the translated text
    contains CJK characters, so viewers render the correct glyphs.
    """
    runs = paragraph.runs
    if not runs:
        return

    runs[0].text = new_text
    for run in runs[1:]:
        run.text = ""

    script = _detect_cjk_script(new_text)
    if script:
        _apply_east_asia_font(runs[0], _SCRIPT_EAST_ASIA_FONT[script])


def build_translated_docx(
    original_bytes: bytes,
    translated_segments: list[dict[str, Any]],
) -> bytes:
    """
    Open the *original* DOCX, walk the same paragraphs / tables / headers
    that the parser visited, and swap in the translated text from
    *translated_segments*.

    Returns the new DOCX as raw bytes.
    """
    doc = Document(io.BytesIO(original_bytes))

    # Build a quick lookup:  segment-id  →  translated text
    lookup: dict[str, str] = {seg["id"]: seg["text"] for seg in translated_segments}

    seg_id = 0

    # ── Body paragraphs ─────────────────────────────────────────
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        key = f"body-{seg_id}"
        if key in lookup:
            _replace_paragraph_text(para, lookup[key])
        seg_id += 1

    # ── Tables ───────────────────────────────────────────────────
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if not cell_text:
                    continue
                key = f"table-{table_idx}-{row_idx}-{col_idx}"
                if key in lookup:
                    # A cell can have multiple paragraphs — put all text
                    # in the first paragraph, clear the rest.
                    paras = cell.paragraphs
                    if paras:
                        _replace_paragraph_text(paras[0], lookup[key])
                        for p in paras[1:]:
                            _replace_paragraph_text(p, "")

    # ── Headers & Footers ────────────────────────────────────────
    hf_seg_id = seg_id  # continue from where body left off (parser did too)
    for section_idx, section in enumerate(doc.sections):
        for hf_type, hf in [("header", section.header), ("footer", section.footer)]:
            if not hf.is_linked_to_previous:
                for para in hf.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    key = f"{hf_type}-{section_idx}-{hf_seg_id}"
                    if key in lookup:
                        _replace_paragraph_text(para, lookup[key])
                    hf_seg_id += 1

    # ── Write out ────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
