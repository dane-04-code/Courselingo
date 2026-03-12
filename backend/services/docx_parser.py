"""DOCX parsing service — extracts translatable text segments from Word documents."""

from __future__ import annotations

from typing import Any
import io

from docx import Document


def extract_text_segments(docx_bytes: bytes) -> list[dict[str, Any]]:
    """
    Open a DOCX from raw bytes and return a flat list of translatable text
    segments.  Each segment is a dict with:

        {
            "id": <unique str>,
            "location": "body" | "table" | "header" | "footer",
            "text": <full paragraph text>,
        }

    Segments with empty / whitespace-only text are skipped.
    """
    doc = Document(io.BytesIO(docx_bytes))
    segments: list[dict[str, Any]] = []
    seg_id = 0

    # ── Body paragraphs ─────────────────────────────────────────
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        segments.append({"id": f"body-{seg_id}", "location": "body", "text": text})
        seg_id += 1

    # ── Tables (each cell contains paragraphs) ──────────────────
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                if not cell_text:
                    continue
                segments.append({
                    "id": f"table-{table_idx}-{row_idx}-{col_idx}",
                    "location": "table",
                    "text": cell_text,
                })
                seg_id += 1

    # ── Headers & Footers ────────────────────────────────────────
    for section_idx, section in enumerate(doc.sections):
        for hf_type, hf in [("header", section.header), ("footer", section.footer)]:
            if not hf.is_linked_to_previous:
                for para in hf.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    segments.append({
                        "id": f"{hf_type}-{section_idx}-{seg_id}",
                        "location": hf_type,
                        "text": text,
                    })
                    seg_id += 1

    return segments
